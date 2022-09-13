import docx
import pandas as pd
import sys, os
import numpy as np
from docxtpl import DocxTemplate

TEMPLATE_FILENAME = "TEMP.docx"

OWNER_STR = 'owner'
NUMBER_STR = 'num'
SQUARE_STR = 'sq'
PART_STR = 'part'
CADASTR_STR = 'cadastr'
FORMALS_STR = 'formals'

def test():
    document = docx.Document("test.docx")

    highlights = []
    for paragraph in document.paragraphs:
        highlight = ""
        for run in paragraph.runs:
            if run.font.highlight_color:
                highlight += run.text
        if highlight:
            highlights.append(highlight)

    for table in document.tables:
        highlight = ""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    highlight = ""
                    for run in paragraph.runs:
                        if run.font.highlight_color:
                            highlight += run.text
                    if highlight:
                        highlights.append(highlight)

    table = document.tables[0]
    cells = table.add_row().cells
    for i in range(len(cells)):
        # вставляем данные в ячейки
        cells[i].text = str(i)

    df1 = pd.read_excel(open('test.xlsx', 'rb'), sheet_name='Sheet1')
    df2 = pd.read_excel(open('test.xlsx', 'rb'), sheet_name='Sheet2')
    print(df1)
    print(df2)
    document.save("test_res.docx")
    doc = DocxTemplate("test.docx")
    context = {'director': "И.И.Иванов"}
    doc.render(context)
    doc.save("result.docx")


def main():
    pass


# Задание: - заменить в бланке для голосования отмеченные желтым элементы на паттерны (теги), которые понимает
# библиотека docxtpl; - написать скрипт на Python, который получает на вход путь к файлу с бланком для голосования,
# путь к файлу с реестром участников, путь к папке для результатов; - после запуска скрипта в папке для результатов
# появляются новые файлы в формате Word, каждый из которых является копией бланка для голосования, где в шапке
# указана информация об одном из участников; названия файлов должны быть в формате: номер помещения, тире, фамилия,
# пробел, имя, пробел, отчетство участника, расширение ".docx"; - для одного участника голосования должен быть только
# один бланк (файл); - после завершения работы скрипта в папке для результатов содержатся именные бланки для всех
# участников голосования.
def openInput(args):
    try:
        if len(args) < 4:
            raise Exception("Not enough input values")
        if len(args) > 4:
            raise Exception("Too much input values")

        pd.read_excel(args[1])
        docx.Document(args[2])
        if not os.path.exists(args[3]):
            os.mkdir(args[3])
            print("Folder '{}' created ".format(args[3]))
        return True
    except Exception as e:
        print(e)
        return False


def createTemplate(docxname, size):
    doc = docx.Document(docxname)
    table = doc.tables[0]
    template_row = [NUMBER_STR, SQUARE_STR, PART_STR, CADASTR_STR, FORMALS_STR]
    cells = table.rows[1].cells

    # all rows
    for j, cell in enumerate(cells):
        text =  '{{' + template_row[j] + str(0) + '}}'
        for i in range(1, size):
            text += ', {{' + template_row[j] + str(i) + '}}'
        cell.text = text

    doc.save(TEMPLATE_FILENAME)
    return


# def createTemplate(docxname, size):
#     doc = docx.Document(docxname)
#     table = doc.tables[0]
#     template_row = [NUMBER_STR, SQUARE_STR, PART_STR, CADASTR_STR, FORMALS_STR]
#     cells = table.rows[1].cells
#
#     # first row
#     for j in range(len(cells)):
#         cells[j].text = '{{' + template_row[j] + str(0) + '}}'
#
#     # next rows
#     for i in range(1, size):
#         cells = table.add_row().cells
#         for j in range(len(cells)):
#             cells[j].text = '{{' + template_row[j] + str(i) + '}}'
#
#     doc.save(TEMPLATE_FILENAME)
#     return


def generateDocxOnTemplate(data: list, path):
    # data
    #  surname name patr num_of_flat space part cadastr formals
    #     0      1    2       3        4     5     6       7
    doc = DocxTemplate(TEMPLATE_FILENAME)

    template_row = [NUMBER_STR, SQUARE_STR, PART_STR, CADASTR_STR, FORMALS_STR]

    context = {OWNER_STR: data[0][0] + ' ' + data[0][1] + ' ' + data[0][2]}
    for i, row in enumerate(data):
        for j, value in enumerate(row[3:]):
            context[template_row[j] + str(i)] = value

    doc.render(context)
    # указана информация об одном из участников; названия файлов должны быть в формате: номер помещения, тире, фамилия,
    # пробел, имя, пробел, отчетство участника, расширение ".docx";
    doc.save(path + '\\' + data[0][3] + '-' + data[0][0] + ' ' + data[0][1] + ' ' + data[0][2] + '.docx')
    return


def getdf(table1: np.ndarray):
    # table1
    # Cтатус | КАДНО | № помещения | Площадь помещ. | Числитель доли | Знаменатель доли | Фамилия / Название ЮЛ
    #     0      1          2              3               4                  5                    6
    # | Имя / ИНН ЮЛ | Отчество / ОГРН ЮЛ | Тип Собственника | № запроса, № выписки | Дата запроса | Вид права
    #        7                   8                   9                  10                 11            12
    # | № государственной регистрации права | Дата госрегистрации | Представитель | Кол-во голосов, кв.м | Доля голосов,%
    #                    13                             14               15                  16                 17

    # needed
    # surname name patr num_of_flat space part_up part_down cadastr formals
    #    0      1    2       3        4      5        6        7       8
    needed = table1[:, [6, 7, 8, 2, 3, 4, 5, 1, 13]]

    # df
    # fullname surname name patr num_of_flat space part cadastr formals
    #     0       1      2    3       4        5     6     7       8

    df = np.array([[str(row[0]).strip() + str(row[1]).strip() + str(row[2]).strip(),  # fullname
                    str(row[0]).strip(),  # surname
                    str(row[1]).strip(),  # name
                    str(row[2]).strip(),  # patr
                    row[3],  # num_of_flat
                    row[4],  # space
                    str(row[5]) + '/' + str(row[6]),  # part
                    row[7],  # cadastr
                    row[8],  # formals
                    ] for row in needed])

    df = df[df[:, 0].argsort()]
    return df


def fillDocs(xlsxname, docxname, path):
    try:
        # Sheet1
        xlsx1 = pd.read_excel(open(xlsxname, 'rb'), sheet_name='Sheet1').fillna('').to_numpy()
        # Sheet2
        xlsx2 = pd.read_excel(open(xlsxname, 'rb'), sheet_name='Sheet2').fillna('').to_numpy()
        # Whole insides of docx as an object to work with
        doc = docx.Document(docxname)

        df = getdf(xlsx1)
        fullname = df[0, 0]  # init
        # table_data = df[0, 1:]  # We dont need fullname
        table_data = [df[0, 1:].tolist()]  # We dont need fullname
        for row in df[1:]:
            if fullname == row[0]:
                table_data.append(row[1:])  # We dont need fullname
            else:
                createTemplate(docxname, len(table_data))
                generateDocxOnTemplate(table_data, path)
                fullname = row[0]
                table_data = [row[1:].tolist()]  # We dont need fullname

        os.remove(TEMPLATE_FILENAME)
        print("DOne!")
    except Exception as e:
        print(e)
        return


if __name__ == "__main__":
    if openInput(sys.argv):
        fillDocs(sys.argv[1], sys.argv[2], sys.argv[3])
